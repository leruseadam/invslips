import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

logger = logging.getLogger(__name__)

class DocumentHandler:
    def __init__(self):
        self.doc = None
        
    def create_document(self, template_path):
        """Create document from template"""
        if not os.path.exists(template_path):
            raise ValueError(f"Template not found: {template_path}")
        self.doc = Document(template_path)
        return self.doc

    def add_content_to_table(self, records):
        """Add content to document replacing placeholders"""
        if not records or not isinstance(records, list):
            return False

        try:
            # Process records in chunks of 4 (for template layout)
            for chunk_index, i in enumerate(range(0, len(records), 4)):
                chunk = records[i:i + 4]
                
                if chunk_index > 0:
                    # Add page break between chunks
                    self.doc.add_page_break()
                
                # Find all paragraphs and tables in the document
                all_paragraphs = []
                for paragraph in self.doc.paragraphs:
                    all_paragraphs.append(paragraph)
                for table in self.doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                all_paragraphs.append(paragraph)

                # Replace placeholders for each record in chunk
                for idx, record in enumerate(chunk, 1):
                    replacements = {
                        f'{{{{Label{idx}.AcceptedDate}}}}': record.get('Accepted Date', ''),
                        f'{{{{Label{idx}.Vendor}}}}': record.get('Vendor', 'Unknown Vendor'),
                        f'{{{{Label{idx}.ProductName}}}}': record.get('Product Name*', ''),
                        f'{{{{Label{idx}.Barcode}}}}': record.get('Barcode*', ''),
                        f'{{{{Label{idx}.QuantityReceived}}}}': str(record.get('Quantity Received*', ''))
                    }
                    
                    # Apply replacements in all paragraphs
                    for paragraph in all_paragraphs:
                        for run in paragraph.runs:
                            for old_text, new_text in replacements.items():
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, str(new_text))
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)

                # Clean up unused placeholders for this chunk
                for idx in range(len(chunk) + 1, 5):
                    empty_replacements = {
                        f'{{{{Label{idx}.AcceptedDate}}}}': '',
                        f'{{{{Label{idx}.Vendor}}}}': '',
                        f'{{{{Label{idx}.ProductName}}}}': '',
                        f'{{{{Label{idx}.Barcode}}}}': '',
                        f'{{{{Label{idx}.QuantityReceived}}}}': ''
                    }
                    
                    for paragraph in all_paragraphs:
                        for run in paragraph.runs:
                            for old_text, new_text in empty_replacements.items():
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, '')

            return True

        except Exception as e:
            logger.error(f"Failed to add content: {str(e)}")
            return False

    def save_document(self, filepath):
        """Save document"""
        try:
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            self.doc.save(filepath)
            return True
        except Exception as e:
            logger.error(f"Failed to save document: {str(e)}")
            return False