import os
import sys
import json
from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer
from io import BytesIO
from docxtpl import DocxTemplate

def chunk_records(records, chunk_size=4):
    """Split records into chunks of specified size"""
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

def adjust_table_font_sizes(doc_path):
    """
    Post-process a DOCX file to dynamically adjust font size inside table cells based on thresholds.
    """
    thresholds = [
        (30, 12),   # <=30 chars → 12pt
        (45, 10),   # <=45 chars → 10pt
        (60, 8),    # <=60 chars → 8pt
        (float('inf'), 7)  # >60 chars → 7pt
    ]

    def get_font_size(text_len):
        for limit, size in thresholds:
            if text_len <= limit:
                return size
        return 7  # Fallback

    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue

                    # If line is Product Name (first line), force 10pt
                    if paragraph == cell.paragraphs[0]:
                        font_size = 10
                    else:
                        font_size = get_font_size(len(text))

                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)

    doc.save(doc_path)

def open_file(path):
    """Open a file using the system's default application"""
    try:
        if sys.platform == "darwin":
            os.system(f'open "{path}"')
        elif sys.platform == "win32":
            os.startfile(path)
        else:
            os.system(f'xdg-open "{path}"')
    except Exception as e:
        print(f"Error opening file: {e}")

def format_json_text(text):
    """Format JSON text for better readability"""
    try:
        if not text.strip():
            return text
        
        parsed = json.loads(text)
        return json.dumps(parsed, indent=2)
    except json.JSONDecodeError:
        return text
    except Exception:
        return text

def run_full_process_inventory_slips(selected_df, config, status_callback=None, progress_callback=None):
    """Process and generate inventory slips"""
    if selected_df.empty:
        if status_callback:
            status_callback("Error: No data selected.")
        return False, "No data selected."

    try:
        # Get settings from config
        items_per_page = int(config['SETTINGS'].get('items_per_page', '4'))
        template_path = config['PATHS'].get('template_path')
        output_dir = config['PATHS'].get('output_dir')
        
        # Ensure output directory exists
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                return False, f"Failed to create output directory: {e}"
        
        if status_callback:
            status_callback("Processing data...")
        
        records = selected_df.to_dict(orient="records")
        pages = []
        
        # Progress calculation
        total_chunks = (len(records) + items_per_page - 1) // items_per_page
        current_chunk = 0
        
        for chunk in chunk_records(records, items_per_page):
            current_chunk += 1
            if progress_callback:
                progress_value = (current_chunk / total_chunks) * 50  # First half of progress
                progress_callback(int(progress_value))
            
            if status_callback:
                status_callback(f"Generating page {current_chunk} of {total_chunks}...")
            
            try:
                tpl = DocxTemplate(template_path)
                context = {}
                
                slot_num = 1
                for rec in chunk:
                    product_name = rec.get("Product Name*", "")
                    barcode = rec.get("Barcode*", "")
                    qty = rec.get("Quantity Received*", rec.get("Quantity*", ""))
                    
                    if not product_name and not barcode and not qty:
                        continue
                    
                    try:
                        qty = int(float(qty))
                    except (ValueError, TypeError):
                        qty = ""
                    
                    context[f"Label{slot_num}"] = {
                        "ProductName": product_name,
                        "Barcode": barcode,
                        "AcceptedDate": rec.get("Accepted Date", ""),
                        "QuantityReceived": qty,
                        "Vendor": rec.get("Vendor", ""),
                        "StrainName": rec.get("Strain Name", ""),
                        "ProductType": rec.get("Product Type*", rec.get("Inventory Type", "")),
                        "THCContent": rec.get("THC Content", ""),
                        "CBDContent": rec.get("CBD Content", "")
                    }
                    slot_num += 1
                
                # Fill empty slots
                for i in range(slot_num, items_per_page + 1):
                    context[f"Label{i}"] = {
                        "ProductName": "",
                        "Barcode": "",
                        "AcceptedDate": "",
                        "QuantityReceived": "",
                        "Vendor": "",
                        "StrainName": "",
                        "ProductType": "",
                        "THCContent": "",
                        "CBDContent": ""
                    }
                
                tpl.render(context)
                buf = BytesIO()
                tpl.save(buf)
                pages.append(Document(buf))
                
            except Exception as e:
                return False, f"Error generating page {current_chunk}: {e}"
        
        if not pages:
            return False, "No documents generated."
        
        if status_callback:
            status_callback("Combining pages...")
        
        master = pages[0]
        composer = Composer(master)
        for i, doc in enumerate(pages[1:]):
            if progress_callback:
                progress_value = 50 + ((i + 1) / len(pages[1:])) * 25  # Second quarter of progress
                progress_callback(int(progress_value))
            composer.append(doc)
        
        now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        outname = f"{now}_inventory_slips.docx"
        outpath = os.path.join(output_dir, outname)
        
        if status_callback:
            status_callback("Saving document...")
        
        master.save(outpath)
        
        if status_callback:
            status_callback("Adjusting formatting...")
        
        adjust_table_font_sizes(outpath)
        
        if progress_callback:
            progress_callback(100)  # Complete progress
        
        if status_callback:
            status_callback(f"Saved to: {outpath}")
        
        # Open file if configured
        auto_open = config['SETTINGS'].getboolean('auto_open', True)
        if auto_open:
            open_file(outpath)
        
        return True, outpath
    
    except Exception as e:
        if status_callback:
            status_callback(f"Error: {e}")
        return False, str(e) 