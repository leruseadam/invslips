import tkinter as tk
from tkinter import ttk, filedialog, messagebox, font
import os
import sys
import json
import datetime
import urllib.request
import pandas as pd
from io import BytesIO
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer
import threading
import configparser
import webbrowser
import re

# Constants
CONFIG_FILE = os.path.expanduser("~/inventory_generator_config.ini")
DEFAULT_SAVE_DIR = os.path.expanduser("~/Downloads")
APP_VERSION = "2.0.0"

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

# Load configurations or create default
def load_config():
    config = configparser.ConfigParser()
    
    # Default configurations
    config['PATHS'] = {
        'template_path': resource_path("templates/InventorySlips.docx"),
        'output_dir': DEFAULT_SAVE_DIR,
        'recent_files': '',
        'recent_urls': ''
    }
    
    config['SETTINGS'] = {
        'items_per_page': '4',
        'auto_open': 'true',
        'theme': 'dark',
        'font_size': '12'
    }
    
    # Load existing config if it exists
    if os.path.exists(CONFIG_FILE):
        config.read(CONFIG_FILE)
    else:
        # Create config file with defaults
        with open(CONFIG_FILE, 'w') as f:
            config.write(f)
    
    return config

def save_config(config):
    with open(CONFIG_FILE, 'w') as f:
        config.write(f)

# Helper to adjust font sizes after rendering
def adjust_table_font_sizes(doc_path):
    """
    Post-process a DOCX file to dynamically adjust font size inside table cells based on thresholds.
    """
    thresholds = [
        (30, 12),   # <=30 chars → 12pt
        (45, 10),   # <=45 chars → 10pt
        (60, 8),    # <=60 chars → 8pt
        (float('inf'), 7)  # >60 chars → 7pt
    ]

    def get_font_size(text_len):
        for limit, size in thresholds:
            if text_len <= limit:
                return size
        return 7  # Fallback

    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue

                    # If line is Product Name (first line), force 10pt
                    if paragraph == cell.paragraphs[0]:
                        font_size = 10
                    else:
                        font_size = get_font_size(len(text))

                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)

    doc.save(doc_path)

# Open files after saving
def open_file(path):
    try:
        if sys.platform == "darwin":
            os.system(f'open "{path}"')
        elif sys.platform == "win32":
            os.startfile(path)
        else:
            os.system(f'xdg-open "{path}"')
    except Exception as e:
        print(f"Error opening file: {e}")

# Split records into chunks
def chunk_records(records, chunk_size=4):
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

# Create tooltips for UI elements
class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)
    
    def show_tooltip(self, event=None):
        x, y, _, _ = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25
        
        self.tooltip = tk.Toplevel(self.widget)
        self.tooltip.wm_overrideredirect(True)
        self.tooltip.wm_geometry(f"+{x}+{y}")
        
        label = tk.Label(
            self.tooltip, 
            text=self.text, 
            justify='left',
            background="#ffffe0", 
            relief="solid", 
            borderwidth=1,
            font=("Arial", 10, "normal")
        )
        label.pack(padx=3, pady=3)
    
    def hide_tooltip(self, event=None):
        if self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None

# Process and save inventory slips - with progress feedback
def run_full_process_inventory_slips(selected_df, config, status_callback=None, progress_callback=None):
    if selected_df.empty:
        if status_callback:
            status_callback("Error: No data selected.")
        return False, "No data selected."

    try:
        # Get settings from config
        items_per_page = int(config['SETTINGS'].get('items_per_page', '4'))
        template_path = config['PATHS'].get('template_path', resource_path("templates/InventorySlips.docx"))
        output_dir = config['PATHS'].get('output_dir', DEFAULT_SAVE_DIR)
        
        # Ensure output directory exists
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                return False, f"Failed to create output directory: {e}"
        
        if status_callback:
            status_callback("Processing data...")
        
        records = selected_df.to_dict(orient="records")
        pages = []
        
        # Progress calculation
        total_chunks = (len(records) + items_per_page - 1) // items_per_page
        current_chunk = 0
        
        for chunk in chunk_records(records, items_per_page):
            current_chunk += 1
            if progress_callback:
                progress_value = (current_chunk / total_chunks) * 50  # First half of progress
                progress_callback(int(progress_value))
            
            if status_callback:
                status_callback(f"Generating page {current_chunk} of {total_chunks}...")
            
            try:
                tpl = DocxTemplate(template_path)
                context = {}
                
                slot_num = 1
                for rec in chunk:
                    product_name = rec.get("Product Name*", "")
                    barcode = rec.get("Barcode*", "")
                    qty = rec.get("Quantity Received*", rec.get("Quantity*", ""))
                    
                    if not product_name and not barcode and not qty:
                        continue
                    
                    try:
                        qty = int(float(qty))
                    except (ValueError, TypeError):
                        qty = ""
                    
                    context[f"Label{slot_num}"] = {
                        "ProductName": product_name,
                        "Barcode": barcode,
                        "AcceptedDate": rec.get("Accepted Date", ""),
                        "QuantityReceived": qty,
                        "Vendor": rec.get("Vendor", ""),
                        "StrainName": rec.get("Strain Name", ""),
                        "ProductType": rec.get("Product Type*", rec.get("Inventory Type", "")),
                        "THCContent": rec.get("THC Content", ""),
                        "CBDContent": rec.get("CBD Content", "")
                    }
                    slot_num += 1
                
                # Fill empty slots
                for i in range(slot_num, items_per_page + 1):
                    context[f"Label{i}"] = {
                        "ProductName": "",
                        "Barcode": "",
                        "AcceptedDate": "",
                        "QuantityReceived": "",
                        "Vendor": "",
                        "StrainName": "",
                        "ProductType": "",
                        "THCContent": "",
                        "CBDContent": ""
                    }
                
                tpl.render(context)
                buf = BytesIO()
                tpl.save(buf)
                pages.append(Document(buf))
                
            except Exception as e:
                return False, f"Error generating page {current_chunk}: {e}"
        
        if not pages:
            return False, "No documents generated."
        
        if status_callback:
            status_callback("Combining pages...")
        
        master = pages[0]
        composer = Composer(master)
        for i, doc in enumerate(pages[1:]):
            if progress_callback:
                progress_value = 50 + ((i + 1) / len(pages[1:])) * 25  # Second quarter of progress
                progress_callback(int(progress_value))
            composer.append(doc)
        
        now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        outname = f"{now}_inventory_slips.docx"
        outpath = os.path.join(output_dir, outname)
        
        if status_callback:
            status_callback("Saving document...")
        
        master.save(outpath)
        
        if status_callback:
            status_callback("Adjusting formatting...")
        
        adjust_table_font_sizes(outpath)
        
        if progress_callback:
            progress_callback(100)  # Complete progress
        
        if status_callback:
            status_callback(f"Saved to: {outpath}")
        
        # Open file if configured
        auto_open = config['SETTINGS'].getboolean('auto_open', True)
        if auto_open:
            open_file(outpath)
        
        return True, outpath
    
    except Exception as e:
        if status_callback:
            status_callback(f"Error: {e}")
        return False, str(e)

# Theme colors
class ThemeColors:
    def __init__(self, theme="dark"):
        # Define theme colors
        self.themes = {
            "dark": {
                "bg_main": "#1E1E2E",
                "bg_secondary": "#181825",
                "fg_main": "#CDD6F4",
                "fg_secondary": "#BAC2DE",
                "accent": "#89B4FA",
                "highlight": "#F5C2E7",
                "button_bg": "#313244",
                "button_fg": "#CDD6F4",
                "checkbox_bg": "#45475A",
                "checkbox_fg": "#F5C2E7",
                "entry_bg": "#313244",
                "entry_fg": "#CDD6F4",
                "success": "#A6E3A1",
                "error": "#F38BA8",
                "warning": "#FAB387"
            },
            "light": {
                "bg_main": "#EFF1F5",
                "bg_secondary": "#CCD0DA",
                "fg_main": "#4C4F69",
                "fg_secondary": "#5C5F77",
                "accent": "#1E66F5",
                "highlight": "#EA76CB",
                "button_bg": "#DCE0E8",
                "button_fg": "#4C4F69",
                "checkbox_bg": "#BCC0CC",
                "checkbox_fg": "#EA76CB",
                "entry_bg": "#DCE0E8",
                "entry_fg": "#4C4F69",
                "success": "#40A02B",
                "error": "#D20F39",
                "warning": "#FE640B"
            },
            "green": {
                "bg_main": "#1A2F1A",
                "bg_secondary": "#132613",
                "fg_main": "#B8E6B8",
                "fg_secondary": "#99CC99",
                "accent": "#40A02B",
                "highlight": "#73D35F",
                "button_bg": "#2D4B2D",
                "button_fg": "#B8E6B8",
                "checkbox_bg": "#3A5F3A",
                "checkbox_fg": "#73D35F",
                "entry_bg": "#2D4B2D",
                "entry_fg": "#B8E6B8",
                "success": "#40A02B",
                "error": "#E64545",
                "warning": "#FFA500"
            }
        }
        
        # Default to dark theme if requested theme doesn't exist
        self.current = self.themes.get(theme, self.themes["dark"])
    
    def get(self, color_name):
        return self.current.get(color_name, "#FFFFFF")
    
    def switch_theme(self, theme_name):
        if theme_name in self.themes:
            self.current = self.themes[theme_name]
            return True
        return False

# Parse Bamboo transfer schema JSON
def parse_bamboo_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Get vendor information
        from_license_number = json_data.get("from_license_number", "")
        from_license_name = json_data.get("from_license_name", "")
        vendor_meta = f"{from_license_number} - {from_license_name}"
        
        # Get transfer date
        raw_date = json_data.get("est_arrival_at", "") or json_data.get("transferred_at", "")
        accepted_date = raw_date.split("T")[0] if "T" in raw_date else raw_date
        
        # Process inventory items
        items = json_data.get("inventory_transfer_items", [])
        records = []
        
        for item in items:
            # Extract THC and CBD content from lab_result_data if available
            thc_content = ""
            cbd_content = ""
            
            lab_data = item.get("lab_result_data", {})
            if lab_data and "potency" in lab_data:
                for potency_item in lab_data["potency"]:
                    if potency_item.get("type") == "total-thc":
                        thc_content = f"{potency_item.get('value', '')}%"
                    elif potency_item.get("type") == "total-cbd":
                        cbd_content = f"{potency_item.get('value', '')}%"
            
            records.append({
                "Product Name*": item.get("product_name", ""),
                "Product Type*": item.get("inventory_type", ""),
                "Quantity Received*": item.get("qty", ""),
                "Barcode*": item.get("inventory_id", "") or item.get("external_id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": item.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Bamboo"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Bamboo transfer data: {e}")

# Parse Cultivera JSON
def parse_cultivera_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Check if Cultivera format
        if not json_data.get("data") or not isinstance(json_data.get("data"), dict):
            raise ValueError("Not a valid Cultivera format")
        
        data = json_data.get("data", {})
        manifest = data.get("manifest", {})
        
        # Get vendor information
        from_license = manifest.get("from_license", {})
        vendor_name = from_license.get("name", "")
        vendor_license = from_license.get("license_number", "")
        vendor_meta = f"{vendor_license} - {vendor_name}" if vendor_license and vendor_name else "Unknown Vendor"
        
        # Get transfer date
        created_at = manifest.get("created_at", "")
        accepted_date = created_at.split("T")[0] if "T" in created_at else created_at
        
        # Process inventory items
        items = manifest.get("items", [])
        records = []
        
        for item in items:
            # Extract product info
            product = item.get("product", {})
            
            # Extract THC and CBD content
            thc_content = ""
            cbd_content = ""
            
            test_results = item.get("test_results", [])
            if test_results:
                for result in test_results:
                    if "thc" in result.get("type", "").lower():
                        thc_content = f"{result.get('percentage', '')}%"
                    elif "cbd" in result.get("type", "").lower():
                        cbd_content = f"{result.get('percentage', '')}%"
            
            records.append({
                "Product Name*": product.get("name", ""),
                "Product Type*": product.get("category", ""),
                "Quantity Received*": item.get("quantity", ""),
                "Barcode*": item.get("barcode", "") or item.get("id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": product.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Cultivera"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Cultivera data: {e}")

# Detect and parse JSON from multiple systems
def parse_inventory_json(json_data):
    """
    Detects the JSON format and parses it accordingly
    """
    if not json_data:
        return None, "No data provided"
    
    try:
        # If data is a string, parse it to JSON
        if isinstance(json_data, str):
            json_data = json.loads(json_data)
        
        # Try parsing as Bamboo
        if "inventory_transfer_items" in json_data:
            return parse_bamboo_data(json_data), "Bamboo"
        
        # Try parsing as Cultivera
        elif "data" in json_data and isinstance(json_data["data"], dict) and "manifest" in json_data["data"]:
            return parse_cultivera_data(json_data), "Cultivera"
        
        # Unknown format
        else:
            return None, "Unknown JSON format. Please use Bamboo or Cultivera format."
    
    except json.JSONDecodeError:
        return None, "Invalid JSON data. Please check the format."
    except Exception as e:
        return None, f"Error parsing data: {str(e)}"

# Main Application Class
class InventorySlipGenerator:
    def __init__(self, root):
        self.root = root
        self.config = load_config()
        self.theme_name = self.config['SETTINGS'].get('theme', 'dark')
        self.colors = ThemeColors(self.theme_name)
        
        self.df = pd.DataFrame()  # Initialize empty DataFrame
        
        self.init_ui()
        self.recent_files = self.config['PATHS'].get('recent_files', '').split('|')
        self.recent_urls = self.config['PATHS'].get('recent_urls', '').split('|')
        self.recent_files = [f for f in self.recent_files if f and os.path.exists(f)]
        self.recent_urls = [u for u in self.recent_urls if u]
        
        self.update_recent_menu()
        
        # Bind window close event
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        
        # Set icon if available
        try:
            if sys.platform == "win32":
                self.root.iconbitmap(resource_path("assets/icon.ico"))
        except:
            pass
    
    def init_ui(self):
        # Configure root window
        self.root.title("Inventory Slip Generator v" + APP_VERSION)
        self.root.configure(bg=self.colors.get("bg_main"))
        
        # Set default size
        width = 850
        height = 750
        self.root.geometry(f"{width}x{height}")
        
        # Create styles for ttk widgets
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Configure ttk styles
        self.style.configure("TButton", 
                             background=self.colors.get("button_bg"),
                             foreground=self.colors.get("button_fg"),
                             font=("Arial", 11))
        
        self.style.configure("TCheckbutton",
                             background=self.colors.get("bg_main"),
                             foreground=self.colors.get("fg_main"),
                             font=("Arial", 11))
        
        self.style.configure("TLabel",
                             background=self.colors.get("bg_main"),
                             foreground=self.colors.get("fg_main"),
                             font=("Arial", 11))
        
        self.style.configure("TFrame",
                             background=self.colors.get("bg_main"))
        
        # Create menu
        self.create_menu()
        
        # Main frame
        self.main_frame = ttk.Frame(self.root, style="TFrame", padding=(20, 10))
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title and Header
        title_frame = ttk.Frame(self.main_frame, style="TFrame")
        title_frame.pack(fill=tk.X, pady=(0, 20))
        
        self.title_label = ttk.Label(
            title_frame, 
            text="Inventory Slip Generator",
            font=("Arial", 18, "bold"),
            style="TLabel"
        )
        self.title_label.pack(side=tk.LEFT)
        
        # Settings button
        self.settings_btn = ttk.Button(
            title_frame,
            text="⚙️ Settings",
            command=self.show_settings,
            style="TButton"
        )
        self.settings_btn.pack(side=tk.RIGHT)
        
        # Create notebook/tabbed interface
        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Create tabs
        self.create_data_tab()
        self.create_preview_tab()
        self.create_bamboo_tab()
        
        # Status frame
        self.status_frame = ttk.Frame(self.root, style="TFrame")
        self.status_frame.pack(fill=tk.X, side=tk.BOTTOM)
        
        # Status bar
        self.status_var = tk.StringVar(value="Ready")
        self.status_label = ttk.Label(
            self.status_frame,
            textvariable=self.status_var,
            style="TLabel",
            padding=(10, 5)
        )
        self.status_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Progress bar
        self.progress_var = tk.IntVar(value=0)
        self.progress_bar = ttk.Progressbar(
            self.status_frame,
            variable=self.progress_var,
            mode='determinate',
            length=200
        )
        self.progress_bar.pack(side=tk.RIGHT, padx=10, pady=5)
    
    def create_menu(self):
        # Main menu bar
        self.menu_bar = tk.Menu(self.root, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        self.root.config(menu=self.menu_bar)
        
        # File menu
        self.file_menu = tk.Menu(self.menu_bar, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        self.menu_bar.add_cascade(label="File", menu=self.file_menu)
        self.file_menu.add_command(label="Open CSV File...", command=self.load_csv)
        self.file_menu.add_command(label="Load from URL...", command=self.show_url_dialog)
        self.file_menu.add_command(label="Paste JSON Data...", command=self.show_json_paste_dialog)

        # Recent files submenu
        self.recent_menu = tk.Menu(self.file_menu, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        self.file_menu.add_cascade(label="Recent Files", menu=self.recent_menu)

        self.file_menu.add_separator()
        self.file_menu.add_command(label="Generate Slips", command=self.on_generate)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="API Settings...", command=self.show_api_settings)
        self.file_menu.add_command(label="Exit", command=self.on_close)
        
        # Edit menu
        self.edit_menu = tk.Menu(self.menu_bar, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        self.menu_bar.add_cascade(label="Edit", menu=self.edit_menu)
        self.edit_menu.add_command(label="Select All", command=lambda: self.select_all_var.set(True) or self.toggle_all())
        self.edit_menu.add_command(label="Deselect All", command=lambda: self.select_all_var.set(False) or self.toggle_all())
        self.edit_menu.add_separator()
        self.edit_menu.add_command(label="Find...", command=self.show_find_dialog)
        
        # View menu
        self.view_menu = tk.Menu(self.menu_bar, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        self.menu_bar.add_cascade(label="View", menu=self.view_menu)
        
        # Theme submenu
        self.theme_menu = tk.Menu(self.view_menu, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        self.view_menu.add_cascade(label="Theme", menu=self.theme_menu)
        self.theme_menu.add_command(label="Dark", command=lambda: self.change_theme("dark"))
        self.theme_menu.add_command(label="Light", command=lambda: self.change_theme("light"))
        self.theme_menu.add_command(label="Green", command=lambda: self.change_theme("green"))
        
        self.view_menu.add_command(label="Refresh Product List", command=self.refresh_product_list)
        
        # Help menu
        self.help_menu = tk.Menu(self.menu_bar, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        self.menu_bar.add_cascade(label="Help", menu=self.help_menu)
        self.help_menu.add_command(label="About", command=self.show_about)
        self.help_menu.add_command(label="Documentation", command=lambda: webbrowser.open("https://github.com/yourusername/inventory-generator/wiki"))
    
    def update_recent_menu(self):
        # Clear the recent menu
        self.recent_menu.delete(0, tk.END)
        
        # Add recent files
        if self.recent_files:
            for file_path in self.recent_files:
                file_name = os.path.basename(file_path)
                self.recent_menu.add_command(
                    label=file_name,
                    command=lambda path=file_path: self.load_csv_from_path(path)
                )
        
        # Add separator if both types exist
        if self.recent_files and self.recent_urls:
            self.recent_menu.add_separator()
        
        # Add recent URLs
        if self.recent_urls:
            for url in self.recent_urls:
                self.recent_menu.add_command(
                    label=url,
                    command=lambda u=url: self.load_from_url(u)
                )
        
        # If no recent items
        if not self.recent_files and not self.recent_urls:
            self.recent_menu.add_command(label="No recent items", state=tk.DISABLED)
    
    
    def create_data_tab(self):
        # Data Source Tab
        self.data_tab = ttk.Frame(self.notebook, style="TFrame")
        self.notebook.add(self.data_tab, text="Data Source")
        
        # URL Entry frame
        url_frame = ttk.Frame(self.data_tab, style="TFrame", padding=(0, 10))
        url_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(url_frame, text="Enter JSON URL:", style="TLabel").pack(anchor=tk.W)
        
        self.url_entry = tk.Entry(
            url_frame,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main")
        )
        # Set default URL
        self.url_entry.insert(0, "https://api-trace.getbamboo.com/shared/manifests/json/ENTER_YOUR_KEY_HERE")
        self.url_entry.pack(fill=tk.X, pady=(5, 0))
        
        # Add right-click menu to URL entry
        self.create_context_menu(self.url_entry)
        
        # Help text below URL entry
        ttk.Label(
            url_frame,
            text="Tip: Use the API Import tab for more options",
            font=("Arial", 9, "italic"),
            style="TLabel"
        ).pack(anchor=tk.W, pady=(5, 0))
        
        # Button frame
        button_frame = ttk.Frame(self.data_tab, style="TFrame")
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # Load buttons
        self.load_json_btn = ttk.Button(
            button_frame,
            text="Load JSON from URL",
            command=self.load_json,
            style="TButton"
        )
        self.load_json_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Label(button_frame, text="Or", style="TLabel").pack(side=tk.LEFT, padx=5)
        
        self.load_csv_btn = ttk.Button(
            button_frame,
            text="Upload CSV File",
            command=self.load_csv,
            style="TButton"
        )
        self.load_csv_btn.pack(side=tk.LEFT, padx=(5, 0))
        
        self.api_import_btn = ttk.Button(
            button_frame,
            text="API Import",
            command=lambda: self.notebook.select(2),  # Switch to API tab
            style="TButton"
        )
        self.api_import_btn.pack(side=tk.LEFT, padx=(10, 0))
        
        # Add tooltips
        ToolTip(self.load_json_btn, "Load inventory data from a JSON URL")
        ToolTip(self.load_csv_btn, "Upload inventory data from a CSV file")
        ToolTip(self.api_import_btn, "Open advanced API import options")
        
        # Selection frame
        select_frame = ttk.Frame(self.data_tab, style="TFrame")
        select_frame.pack(fill=tk.X, padx=10, pady=(20, 10))
        
        ttk.Label(select_frame, text="Select products to include:", style="TLabel").pack(anchor=tk.W)
        
        # Select all checkbox
        self.select_all_var = tk.BooleanVar(value=True)
        self.select_all_cb = ttk.Checkbutton(
            select_frame,
            text="Select / Deselect All",
            variable=self.select_all_var,
            command=self.toggle_all,
            style="TCheckbutton"
        )
        self.select_all_cb.pack(pady=(10, 10), anchor=tk.W)
        
        # Search and filter frame
        filter_frame = ttk.Frame(self.data_tab, style="TFrame")
        filter_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        ttk.Label(filter_frame, text="Search:", style="TLabel").pack(side=tk.LEFT, padx=(0, 5))
        
        self.search_var = tk.StringVar()
        self.search_var.trace_add("write", self.on_search)
        
        self.search_entry = tk.Entry(
            filter_frame,
            textvariable=self.search_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main")
        )
        self.search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Clear search button
        self.clear_search_btn = ttk.Button(
            filter_frame,
            text="×",
            width=2,
            command=lambda: self.search_var.set(""),
            style="TButton"
        )
        self.clear_search_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # Create scrollable frame for products
        self.create_scrollable_product_list()
        
        # Generate button
        self.generate_btn = ttk.Button(
            self.data_tab,
            text="Generate Inventory Slips",
            command=self.on_generate,
            style="TButton"
        )
        self.generate_btn.pack(pady=20)
    
    def create_scrollable_product_list(self):
        # Container frame for the canvas and scrollbar
        self.list_container = ttk.Frame(self.data_tab, style="TFrame")
        self.list_container.pack(fill=tk.BOTH, expand=True, padx=10)
        
        # Canvas for scrolling
        self.canvas = tk.Canvas(
            self.list_container,
            bg=self.colors.get("bg_secondary"),
            highlightthickness=0
        )
        
        # Scrollbar
        self.scrollbar = ttk.Scrollbar(
            self.list_container,
            orient="vertical",
            command=self.canvas.yview
        )
        
        # Frame inside canvas to hold products
        self.product_frame = ttk.Frame(self.canvas, style="TFrame")
        
        # Configure canvas and scrolling
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Create window inside canvas
        self.canvas_window = self.canvas.create_window(
            (0, 0),
            window=self.product_frame,
            anchor="nw",
            tags="product_frame"
        )
        
        # Bind events for scrolling and resizing
        self.product_frame.bind("<Configure>", self.on_frame_configure)
        self.canvas.bind("<Configure>", self.on_canvas_configure)
        
        # Bind mouse wheel for scrolling
        self.canvas.bind_all("<MouseWheel>", self.on_mousewheel)
        self.canvas.bind_all("<Button-4>", self.on_mousewheel)
        self.canvas.bind_all("<Button-5>", self.on_mousewheel)
        
        # Initialize dictionaries for product variables
        self.product_vars = {}
        self.group_vars = {}
        
        # Initially display a message when no products are loaded
        self.empty_label = ttk.Label(
            self.product_frame,
            text="No products loaded. Please load data from CSV or JSON.",
            style="TLabel",
            font=("Arial", 12)
        )
        self.empty_label.pack(pady=50)
    
    def create_preview_tab(self):
        # Preview Tab
        self.preview_tab = ttk.Frame(self.notebook, style="TFrame")
        self.notebook.add(self.preview_tab, text="Preview & Settings")
        
        # Output directory frame
        output_frame = ttk.Frame(self.preview_tab, style="TFrame", padding=(10, 10))
        output_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(output_frame, text="Output Directory:", style="TLabel").pack(anchor=tk.W)
        
        output_path_frame = ttk.Frame(output_frame, style="TFrame")
        output_path_frame.pack(fill=tk.X, pady=(5, 0))
        
        self.output_dir_var = tk.StringVar(value=self.config['PATHS'].get('output_dir', DEFAULT_SAVE_DIR))
        
        self.output_entry = tk.Entry(
            output_path_frame,
            textvariable=self.output_dir_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main")
        )
        self.output_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.browse_btn = ttk.Button(
            output_path_frame,
            text="Browse...",
            command=self.browse_output_dir,
            style="TButton"
        )
        self.browse_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # Items per page frame
        items_frame = ttk.Frame(self.preview_tab, style="TFrame", padding=(10, 10))
        items_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        ttk.Label(items_frame, text="Items per page:", style="TLabel").pack(side=tk.LEFT)
        
        self.items_per_page_var = tk.StringVar(value=self.config['SETTINGS'].get('items_per_page', '4'))
        
        items_combo = ttk.Combobox(
            items_frame,
            textvariable=self.items_per_page_var,
            values=["2", "4", "6", "8"],
            width=5,
            state="readonly"
        )
        items_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        # Auto-open checkbox
        self.auto_open_var = tk.BooleanVar(value=self.config['SETTINGS'].getboolean('auto_open', True))
        auto_open_cb = ttk.Checkbutton(
            items_frame,
            text="Auto-open after generation",
            variable=self.auto_open_var,
            style="TCheckbutton"
        )
        auto_open_cb.pack(side=tk.RIGHT)
        
        # Template frame
        template_frame = ttk.Frame(self.preview_tab, style="TFrame", padding=(10, 10))
        template_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        ttk.Label(template_frame, text="Template File:", style="TLabel").pack(anchor=tk.W)
        
        template_path_frame = ttk.Frame(template_frame, style="TFrame")
        template_path_frame.pack(fill=tk.X, pady=(5, 0))
        
        self.template_path_var = tk.StringVar(value=self.config['PATHS'].get('template_path', resource_path("templates/InventorySlips.docx")))
        
        self.template_entry = tk.Entry(
            template_path_frame,
            textvariable=self.template_path_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main")
        )
        self.template_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.template_browse_btn = ttk.Button(
            template_path_frame,
            text="Browse...",
            command=self.browse_template,
            style="TButton"
        )
        self.template_browse_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # Preview placeholder
        preview_label_frame = ttk.LabelFrame(self.preview_tab, text="Template Preview", style="TFrame")
        preview_label_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.preview_text = tk.Text(
            preview_label_frame,
            font=("Arial", 10),
            bg=self.colors.get("bg_secondary"),
            fg=self.colors.get("fg_main"),
            height=10,
            state=tk.DISABLED
        )
        self.preview_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Apply settings button
        self.apply_settings_btn = ttk.Button(
            self.preview_tab,
            text="Apply Settings",
            command=self.apply_settings,
            style="TButton"
        )
        self.apply_settings_btn.pack(pady=(0, 10))
    
    def create_bamboo_tab(self):
        # API Import Tab
        self.bamboo_tab = ttk.Frame(self.notebook, style="TFrame")
        self.notebook.add(self.bamboo_tab, text="API Import")
        
        # Info section
        info_frame = ttk.Frame(self.bamboo_tab, style="TFrame", padding=(10, 10))
        info_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(
            info_frame,
            text="Import data from Bamboo or Cultivera",
            font=("Arial", 12, "bold"),
            style="TLabel"
        ).pack(anchor=tk.W)
        
        ttk.Label(
            info_frame,
            text="This tab allows you to import data directly from Bamboo or Cultivera formats.\n"
                 "You can paste JSON data into the text area below or load from a file.",
            wraplength=800,
            style="TLabel"
        ).pack(anchor=tk.W, pady=(5, 0))
        
        # API selection frame
        api_selection_frame = ttk.Frame(self.bamboo_tab, style="TFrame")
        api_selection_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        # API selection radiobuttons
        self.api_var = tk.StringVar(value="auto")
        
        ttk.Label(
            api_selection_frame,
            text="API Format:",
            style="TLabel"
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Radiobutton(
            api_selection_frame,
            text="Auto-detect",
            variable=self.api_var,
            value="auto",
            style="TRadiobutton"
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Radiobutton(
            api_selection_frame,
            text="Bamboo",
            variable=self.api_var,
            value="bamboo",
            style="TRadiobutton"
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Radiobutton(
            api_selection_frame,
            text="Cultivera",
            variable=self.api_var,
            value="cultivera",
            style="TRadiobutton"
        ).pack(side=tk.LEFT)
        
        # Button frame
        button_frame = ttk.Frame(self.bamboo_tab, style="TFrame")
        button_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        self.load_bamboo_file_btn = ttk.Button(
            button_frame,
            text="Load JSON File",
            command=self.load_bamboo_file,
            style="TButton"
        )
        self.load_bamboo_file_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.paste_json_btn = ttk.Button(
            button_frame,
            text="Paste JSON Data",
            command=self.show_json_paste_dialog,
            style="TButton"
        )
        self.paste_json_btn.pack(side=tk.LEFT, padx=(5, 0))
        
        self.clear_json_btn = ttk.Button(
            button_frame,
            text="Clear JSON",
            command=self.clear_json_data,
            style="TButton"
        )
        self.clear_json_btn.pack(side=tk.LEFT, padx=(5, 0))
        
        self.fetch_api_btn = ttk.Button(
            button_frame,
            text="Fetch from API",
            command=self.show_api_fetch_dialog,
            style="TButton"
        )
        self.fetch_api_btn.pack(side=tk.LEFT, padx=(5, 0))
        
        # JSON text area with scrollbar
        json_frame = ttk.LabelFrame(self.bamboo_tab, text="JSON Data", style="TFrame")
        json_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        
        # Scrollbar
        json_scroll = ttk.Scrollbar(json_frame)
        json_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Text widget for JSON
        # Text widget for JSON
        self.json_text = tk.Text(
            json_frame,
            font=("Consolas", 10),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main"),
            yscrollcommand=json_scroll.set
        )
        self.json_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        json_scroll.config(command=self.json_text.yview)
        
        # Create context menu for text area
        self.create_context_menu(self.json_text)
        
        # Import button
        self.import_json_btn = ttk.Button(
            self.bamboo_tab,
            text="Import Data",
            command=self.import_bamboo_data,
            style="TButton"
        )
        self.import_json_btn.pack(pady=(0, 10))
        
        # Info text at bottom
        ttk.Label(
            self.bamboo_tab,
            text="Supported formats: Bamboo Transfer Schema and Cultivera Manifest",
            font=("Arial", 9),
            style="TLabel"
        ).pack(pady=(0, 5))
    
    def style_product_row(self, row_frame, source_system):
        """Apply styling to product rows based on their source system"""
        if source_system == "Bamboo":
            row_frame.configure(bg=self.colors.get("success"))
            for child in row_frame.winfo_children():
                if isinstance(child, ttk.Label):
                    child.configure(foreground=self.colors.get("success"))
        elif source_system == "Cultivera":
            row_frame.configure(bg=self.colors.get("accent"))
            for child in row_frame.winfo_children():
                if isinstance(child, ttk.Label):
                    child.configure(foreground=self.colors.get("accent"))
    
    def refresh_product_list(self):
        # Clear existing product list
        for widget in self.product_frame.winfo_children():
            widget.destroy()
        
        self.product_vars.clear()
        self.group_vars.clear()
        
        # If no data, show empty message
        if self.df.empty:
            self.empty_label = ttk.Label(
                self.product_frame,
                text="No products loaded. Please load data from CSV or JSON.",
                style="TLabel",
                font=("Arial", 12)
            )
            self.empty_label.pack(pady=50)
            return
        
        # Add sources legend
        if "Source System" in self.df.columns and len(self.df["Source System"].unique()) > 1:
            legend_frame = ttk.Frame(self.product_frame, style="TFrame")
            legend_frame.pack(fill=tk.X, pady=(5, 10), padx=5)
            
            ttk.Label(
                legend_frame,
                text="Data Sources:",
                style="TLabel",
                font=("Arial", 10, "bold")
            ).pack(side=tk.LEFT, padx=(0, 10))
            
            for source in sorted(self.df["Source System"].unique()):
                source_color = self.colors.get("success") if source == "Bamboo" else self.colors.get("accent")
                
                source_label = ttk.Label(
                    legend_frame,
                    text=source,
                    style="TLabel",
                    foreground=source_color,
                    font=("Arial", 10, "bold")
                )
                source_label.pack(side=tk.LEFT, padx=(0, 20))
        
        # Group data by product type
        grouped = self.df.groupby("Product Type*") if "Product Type*" in self.df.columns else {"All Products": self.df}
        
        # For each product type, create a group
        for product_type, group_df in grouped:
            # Add group select checkbox
            group_var = tk.BooleanVar(value=True)
            self.group_vars[product_type] = group_var
            
            # Create frame for group header
            group_frame = ttk.Frame(self.product_frame, style="TFrame")
            group_frame.pack(fill=tk.X, pady=(10, 5), padx=5)
            
            # Group checkbox
            group_cb = ttk.Checkbutton(
                group_frame,
                text=f"Select/Deselect {product_type}",
                variable=group_var,
                command=lambda pt=product_type: self.toggle_group(pt),
                style="TCheckbutton"
            )
            group_cb.pack(fill=tk.X, pady=2)
            
            # Count items in group
            count_label = ttk.Label(
                group_frame,
                text=f"({len(group_df)} items)",
                style="TLabel",
                font=("Arial", 9)
            )
            count_label.pack(side=tk.RIGHT, padx=(0, 5))
            
            # Add each product under the group
            for idx, row in group_df.iterrows():
                product_name = row.get("Product Name*", "")
                sku = str(row.get("Barcode*", "")).strip()
                strain_name = row.get("Strain Name", "")
                source_system = row.get("Source System", "")
                
                qty = ""
                if "Quantity Received*" in row and pd.notna(row["Quantity Received*"]):
                    qty = row["Quantity Received*"]
                elif "Quantity*" in row and pd.notna(row["Quantity*"]):
                    qty = row["Quantity*"]
                
                if pd.notna(qty):
                    try:
                        qty = int(float(qty))
                    except (ValueError, TypeError):
                        qty = str(qty)
                else:
                    qty = ""
                
                # Create row with product info
                row_frame = ttk.Frame(self.product_frame, style="TFrame")
                row_frame.pack(fill=tk.X, pady=2, padx=10)
                
                # Store metadata for search functionality
                row_frame.product_name = product_name
                row_frame.product_type = product_type
                row_frame.strain_name = strain_name
                row_frame.source_system = source_system
                
                # Checkbox for selection
                var = tk.BooleanVar(value=True)
                cb = ttk.Checkbutton(
                    row_frame,
                    text=f"{product_name}",
                    variable=var,
                    style="TCheckbutton"
                )
                cb.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                # SKU and Quantity labels
                info_frame = ttk.Frame(row_frame, style="TFrame")
                info_frame.pack(side=tk.RIGHT)
                
                # Source system indicator (small colored square)
                if source_system:
                    source_color = self.colors.get("success") if source_system == "Bamboo" else self.colors.get("accent")
                    source_indicator = ttk.Label(
                        info_frame,
                        text=f"• {source_system}",
                        style="TLabel",
                        font=("Arial", 9),
                        foreground=source_color
                    )
                    source_indicator.pack(side=tk.TOP, anchor=tk.E)
                
                if strain_name:
                    ttk.Label(
                        info_frame,
                        text=f"Strain: {strain_name}",
                        style="TLabel",
                        font=("Arial", 9)
                    ).pack(side=tk.TOP, anchor=tk.E)
                
                ttk.Label(
                    info_frame,
                    text=f"SKU: {sku}",
                    style="TLabel",
                    font=("Arial", 9)
                ).pack(side=tk.TOP, anchor=tk.E)
                
                ttk.Label(
                    info_frame,
                    text=f"Qty: {qty}",
                    style="TLabel",
                    font=("Arial", 10, "bold")
                ).pack(side=tk.TOP, anchor=tk.E)
                
                # Add THC/CBD content if available
                thc_content = row.get("THC Content", "")
                cbd_content = row.get("CBD Content", "")
                
                if thc_content or cbd_content:
                    potency_text = []
                    if thc_content:
                        potency_text.append(f"THC: {thc_content}")
                    if cbd_content:
                        potency_text.append(f"CBD: {cbd_content}")
                    
                    ttk.Label(
                        info_frame,
                        text=" | ".join(potency_text),
                        style="TLabel",
                        font=("Arial", 9)
                    ).pack(side=tk.TOP, anchor=tk.E)
                
                # Store variable and product type for bulk selection
                self.product_vars[idx] = (var, product_type)
    
    def on_search(self, *args):
        search_text = self.search_var.get().lower()
        
        # Hide all products initially
        for widget in self.product_frame.winfo_children():
            if isinstance(widget, ttk.Frame) and hasattr(widget, 'product_type'):
                widget.pack_forget()
        
        # If no search text, show all products
        if not search_text:
            for widget in self.product_frame.winfo_children():
                if isinstance(widget, ttk.Frame) and hasattr(widget, 'product_type'):
                    widget.pack(fill=tk.X, pady=2, padx=5)
            return
        
        # Show only matching products
        for widget in self.product_frame.winfo_children():
            if isinstance(widget, ttk.Frame) and hasattr(widget, 'product_type'):
                product_name = getattr(widget, 'product_name', '').lower()
                product_type = getattr(widget, 'product_type', '').lower()
                strain_name = getattr(widget, 'strain_name', '').lower()
                source_system = getattr(widget, 'source_system', '').lower()
                
                if (search_text in product_name or 
                    search_text in product_type or 
                    search_text in strain_name or
                    search_text in source_system):
                    widget.pack(fill=tk.X, pady=2, padx=5)
    
    def import_bamboo_data(self):
        # Get JSON data from text area
        json_data = self.json_text.get(1.0, tk.END)
        if not json_data.strip():
            messagebox.showwarning("Warning", "No JSON data to import.")
            return
        
        # Get selected API format
        api_format = self.api_var.get()
        
        try:
            if api_format == "bamboo":
                # Force Bamboo format
                data = json.loads(json_data)
                self.df = parse_bamboo_data(data)
                format_type = "Bamboo"
            elif api_format == "cultivera":
                # Force Cultivera format
                data = json.loads(json_data)
                self.df = parse_cultivera_data(data)
                format_type = "Cultivera"
            else:
                # Auto-detect format
                self.process_json_data(json_data)
                return
            
            # Update UI for specific format
            self.notebook.tab(2, text=f"{format_type} Import")
            self.status_var.set(f"{format_type} data imported successfully.")
            self.progress_var.set(100)
            
            # Sort the DataFrame for better organization
            if not self.df.empty:
                try:
                    sort_cols = []
                    if "Product Type*" in self.df.columns:
                        sort_cols.append("Product Type*")
                    if "Product Name*" in self.df.columns:
                        sort_cols.append("Product Name*")
                    
                    if sort_cols:
                        self.df = self.df.sort_values(sort_cols, ascending=[True, True])
                except:
                    pass  # If sorting fails, continue without sorting
            
            self.refresh_product_list()
            
            # Switch to Data tab
            self.notebook.select(0)
            
            # Reset progress after a delay
            self.root.after(2000, lambda: self.progress_var.set(0))
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to import data:\n{e}")
            self.status_var.set("Failed to import data.")
            self.progress_var.set(0)
    
    def create_context_menu(self, widget):
        context_menu = tk.Menu(widget, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        context_menu.add_command(label="Cut", command=lambda: widget.event_generate('<<Cut>>'))
        context_menu.add_command(label="Copy", command=lambda: widget.event_generate('<<Copy>>'))
        context_menu.add_command(label="Paste", command=lambda: widget.event_generate('<<Paste>>'))
        context_menu.add_separator()
        context_menu.add_command(label="Select All", command=lambda: widget.event_generate('<<SelectAll>>'))
        
        if isinstance(widget, tk.Text):
            context_menu.add_separator()
            context_menu.add_command(label="Format JSON", command=lambda: self.format_json_text(widget))
        
        def show_context_menu(event):
            try:
                context_menu.tk_popup(event.x_root, event.y_root)
            finally:
                context_menu.grab_release()
        
        widget.bind("<Button-3>", show_context_menu)  # Right-click on Windows/Linux
        if sys.platform == 'darwin':
            widget.bind("<Button-2>", show_context_menu)  # Right-click on macOS
    
    def format_json_text(self, text_widget):
        """Format the JSON in a text widget for better readability"""
        try:
            content = text_widget.get(1.0, tk.END)
            if not content.strip():
                return
            
            parsed = json.loads(content)
            formatted = json.dumps(parsed, indent=2)
            
            text_widget.delete(1.0, tk.END)
            text_widget.insert(tk.END, formatted)
            
            self.status_var.set("JSON formatted successfully.")
        except json.JSONDecodeError:
            messagebox.showerror("Error", "Invalid JSON. Cannot format.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to format JSON:\n{e}")
    
    def on_frame_configure(self, event=None):
        # Update the scrollregion to encompass the inner frame
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def on_canvas_configure(self, event=None):
        # When the canvas is resized, also resize the inner frame
        width = event.width
        self.canvas.itemconfig(self.canvas_window, width=width)
    
    def on_mousewheel(self, event):
        # Respond to mouse wheel events for scrolling
        if sys.platform == 'darwin':
            self.canvas.yview_scroll(-1 * event.delta, "units")
        else:
            if hasattr(event, 'num'):
                if event.num == 4:
                    self.canvas.yview_scroll(-1, "units")
                elif event.num == 5:
                    self.canvas.yview_scroll(1, "units")
            else:
                self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    
    def toggle_all(self):
        select = self.select_all_var.get()
        
        # First: toggle all individual product checkboxes
        for var, _ in self.product_vars.values():
            var.set(select)
        
        # Then: toggle all group (product type) checkboxes
        for group_var in self.group_vars.values():
            group_var.set(select)
    
    def toggle_group(self, product_type):
        group_select = self.group_vars.get(product_type)
        if group_select is None:
            return
        
        select = group_select.get()
        
        for idx, (var, ptype) in self.product_vars.items():
            if ptype == product_type:
                var.set(select)
    
    def show_api_fetch_dialog(self):
        api_dialog = tk.Toplevel(self.root)
        api_dialog.title("Fetch from API")
        api_dialog.geometry("600x200")
        api_dialog.resizable(True, False)
        api_dialog.transient(self.root)
        api_dialog.grab_set()
        
        # Make dialog modal
        api_dialog.focus_set()
        
        # Configure dialog
        api_dialog.configure(bg=self.colors.get("bg_main"))
        
        # API frame
        api_frame = ttk.Frame(api_dialog, style="TFrame", padding=(10, 10))
        api_frame.pack(fill=tk.BOTH, expand=True)
        
        # API type selector
        api_type_frame = ttk.Frame(api_frame, style="TFrame")
        api_type_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(
            api_type_frame,
            text="API Type:",
            style="TLabel"
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        api_type_var = tk.StringVar(value=self.api_var.get())
        api_combo = ttk.Combobox(
            api_type_frame,
            textvariable=api_type_var,
            values=["auto", "bamboo", "cultivera"],
            state="readonly",
            width=10
        )
        api_combo.pack(side=tk.LEFT)
        
        # URL frame
        url_frame = ttk.Frame(api_frame, style="TFrame")
        url_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(
            url_frame,
            text="API URL:",
            style="TLabel"
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        url_var = tk.StringVar()
        url_entry = tk.Entry(
            url_frame,
            textvariable=url_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main"),
            width=50
        )
        url_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        url_entry.focus_set()
        
        # Authentication frame (for future use)
        auth_frame = ttk.Frame(api_frame, style="TFrame")
        auth_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(
            auth_frame,
            text="API Key (optional):",
            style="TLabel"
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        api_key_var = tk.StringVar()
        api_key_entry = tk.Entry(
            auth_frame,
            textvariable=api_key_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main"),
            show="*"
        )
        api_key_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Button frame
        button_frame = ttk.Frame(api_frame, style="TFrame")
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        ttk.Button(
            button_frame,
            text="Fetch",
            command=lambda: self.fetch_from_api(
                url_var.get(),
                api_type_var.get(),
                api_key_var.get(),
                api_dialog
            ),
            style="TButton"
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        ttk.Button(
            button_frame,
            text="Cancel",
            command=api_dialog.destroy,
            style="TButton"
        ).pack(side=tk.RIGHT)
    
    def fetch_from_api(self, url, api_type, api_key=None, dialog=None):
        if not url:
            messagebox.showerror("Error", "Please enter an API URL.")
            return
        
        if dialog:
            dialog.destroy()
        
        self.status_var.set(f"Fetching data from API...")
        self.progress_var.set(10)
        
        headers = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        
        def fetch_data():
            try:
                req = urllib.request.Request(url, headers=headers)
                with urllib.request.urlopen(req) as resp:
                    data = json.loads(resp.read().decode())
                
                # Process data based on format
                self.root.after(0, lambda: self.process_api_data(data, api_type))
                
                # Add to recent URLs if not already there
                if url not in self.recent_urls:
                    self.recent_urls.insert(0, url)
                    self.recent_urls = self.recent_urls[:10]  # Keep only 10 most recent
                    self.config['PATHS']['recent_urls'] = '|'.join(self.recent_urls)
                    save_config(self.config)
                    self.update_recent_menu()
            
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to fetch API data:\n{e}"))
                self.root.after(0, lambda: self.status_var.set("Failed to fetch data."))
                self.root.after(0, lambda: self.progress_var.set(0))
        
        # Run in a separate thread to prevent UI freezing
        threading.Thread(target=fetch_data, daemon=True).start()
    
    def process_api_data(self, data, api_type="auto"):
        self.status_var.set("Processing API data...")
        self.progress_var.set(30)
        
        try:
            # Process based on specified API type
            if api_type == "bamboo":
                result_df = parse_bamboo_data(data)
                format_type = "Bamboo"
            elif api_type == "cultivera":
                result_df = parse_cultivera_data(data)
                format_type = "Cultivera"
            else:  # auto-detect
                result_df, format_type = parse_inventory_json(data)
            
            if result_df is None or result_df.empty:
                messagebox.showerror("Error", f"Could not process {api_type} data.")
                self.status_var.set("Failed to process data.")
                self.progress_var.set(0)
                return
            
            # Success
            self.df = result_df
            
            # Update JSON text area
            self.json_text.delete(1.0, tk.END)
            self.json_text.insert(tk.END, json.dumps(data, indent=2))
            
            # Update tab name to reflect the format
            self.notebook.tab(2, text=f"{format_type} Import")
            
            # Sort the DataFrame for better organization
            if not self.df.empty:
                try:
                    sort_cols = []
                    if "Product Type*" in self.df.columns:
                        sort_cols.append("Product Type*")
                    if "Product Name*" in self.df.columns:
                        sort_cols.append("Product Name*")
                    
                    if sort_cols:
                        self.df = self.df.sort_values(sort_cols, ascending=[True, True])
                except:
                    pass  # If sorting fails, continue without sorting
            
            self.refresh_product_list()
            self.status_var.set(f"{format_type} data processed successfully.")
            self.progress_var.set(100)
            
            # Switch to Data tab
            self.notebook.select(0)
            
            # Reset progress after a delay
            self.root.after(2000, lambda: self.progress_var.set(0))
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process API data:\n{e}")
            self.status_var.set("Failed to process data.")
            self.progress_var.set(0) (padx=5, pady=5)
        json_scroll.config(command=self.json_text.yview)
        
        # Create context menu for text area
        self.create_context_menu(self.json_text)
        
        # Import button
        self.import_json_btn = ttk.Button(
            self.bamboo_tab,
            text="Import Bamboo Data",
            command=self.import_bamboo_data,
            style="TButton"
        )
        self.import_json_btn.pack(pady=(0, 10))
    
    def create_context_menu(self, widget):
        context_menu = tk.Menu(widget, tearoff=0, bg=self.colors.get("bg_secondary"), fg=self.colors.get("fg_main"))
        context_menu.add_command(label="Cut", command=lambda: widget.event_generate('<<Cut>>'))
        context_menu.add_command(label="Copy", command=lambda: widget.event_generate('<<Copy>>'))
        context_menu.add_command(label="Paste", command=lambda: widget.event_generate('<<Paste>>'))
        context_menu.add_separator()
        context_menu.add_command(label="Select All", command=lambda: widget.event_generate('<<SelectAll>>'))
        
        def show_context_menu(event):
            try:
                context_menu.tk_popup(event.x_root, event.y_root)
            finally:
                context_menu.grab_release()
        
        widget.bind("<Button-3>", show_context_menu)  # Right-click on Windows/Linux
        if sys.platform == 'darwin':
            widget.bind("<Button-2>", show_context_menu)  # Right-click on macOS
    
    def on_frame_configure(self, event=None):
        # Update the scrollregion to encompass the inner frame
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def on_canvas_configure(self, event=None):
        # When the canvas is resized, also resize the inner frame
        width = event.width
        self.canvas.itemconfig(self.canvas_window, width=width)
    
    def on_mousewheel(self, event):
        # Respond to mouse wheel events for scrolling
        if sys.platform == 'darwin':
            self.canvas.yview_scroll(-1 * event.delta, "units")
        else:
            if event.num == 4:
                self.canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                self.canvas.yview_scroll(1, "units")
            else:
                self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    
    def toggle_all(self):
        select = self.select_all_var.get()
        
        # First: toggle all individual product checkboxes
        for var, _ in self.product_vars.values():
            var.set(select)
        
        # Then: toggle all group (product type) checkboxes
        for group_var in self.group_vars.values():
            group_var.set(select)
    
    def toggle_group(self, product_type):
        group_select = self.group_vars.get(product_type)
        if group_select is None:
            return
        
        select = group_select.get()
        
        for idx, (var, ptype) in self.product_vars.items():
            if ptype == product_type:
                var.set(select)
    
    def on_search(self, *args):
        search_text = self.search_var.get().lower()
        
        # Hide all products initially
        for widget in self.product_frame.winfo_children():
            if isinstance(widget, tk.Frame) and hasattr(widget, 'product_type'):
                widget.pack_forget()
        
        # If no search text, show all products
        if not search_text:
            for widget in self.product_frame.winfo_children():
                if isinstance(widget, tk.Frame) and hasattr(widget, 'product_type'):
                    widget.pack(fill=tk.X, pady=2, padx=5)
            return
        
        # Show only matching products
        for widget in self.product_frame.winfo_children():
            if isinstance(widget, tk.Frame) and hasattr(widget, 'product_type'):
                product_name = getattr(widget, 'product_name', '').lower()
                product_type = getattr(widget, 'product_type', '').lower()
                strain_name = getattr(widget, 'strain_name', '').lower()
                
                if (search_text in product_name or 
                    search_text in product_type or 
                    search_text in strain_name):
                    widget.pack(fill=tk.X, pady=2, padx=5)
    
    def show_find_dialog(self):
        find_dialog = tk.Toplevel(self.root)
        find_dialog.title("Find Product")
        find_dialog.geometry("400x100")
        find_dialog.resizable(False, False)
        find_dialog.transient(self.root)
        find_dialog.grab_set()
        
        # Make dialog modal
        find_dialog.focus_set()
        
        # Configure find dialog
        find_dialog.configure(bg=self.colors.get("bg_main"))
        
        # Frame for find interface
        find_frame = ttk.Frame(find_dialog, style="TFrame", padding=(10, 10))
        find_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(find_frame, text="Find text:", style="TLabel").pack(anchor=tk.W)
        
        find_var = tk.StringVar()
        find_entry = tk.Entry(
            find_frame,
            textvariable=find_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main")
        )
        find_entry.pack(fill=tk.X, pady=(5, 10))
        find_entry.focus_set()
        
        button_frame = ttk.Frame(find_frame, style="TFrame")
        button_frame.pack(fill=tk.X)
        
        ttk.Button(
            button_frame,
            text="Find",
            command=lambda: self.perform_find(find_var.get(), find_dialog),
            style="TButton"
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        ttk.Button(
            button_frame,
            text="Cancel",
            command=find_dialog.destroy,
            style="TButton"
        ).pack(side=tk.RIGHT)
        
        # Bind Enter key to Find command
        find_entry.bind("<Return>", lambda event: self.perform_find(find_var.get(), find_dialog))
    
    def perform_find(self, search_text, dialog):
        if not search_text:
            dialog.destroy()
            return
        
        self.search_var.set(search_text)
        dialog.destroy()
        
        # Switch to Data tab
        self.notebook.select(0)
    
    def browse_output_dir(self):
        directory = filedialog.askdirectory(
            initialdir=self.output_dir_var.get(),
            title="Select Output Directory"
        )
        
        if directory:
            self.output_dir_var.set(directory)
    
    def browse_template(self):
        file_path = filedialog.askopenfilename(
            initialdir=os.path.dirname(self.template_path_var.get()),
            title="Select Template File",
            filetypes=[("Word Documents", "*.docx")]
        )
        
        if file_path:
            self.template_path_var.set(file_path)
    
    def apply_settings(self):
        # Save settings to config
        self.config['PATHS']['output_dir'] = self.output_dir_var.get()
        self.config['PATHS']['template_path'] = self.template_path_var.get()
        self.config['SETTINGS']['items_per_page'] = self.items_per_page_var.get()
        self.config['SETTINGS']['auto_open'] = str(self.auto_open_var.get())
        
        save_config(self.config)
        
        # Update status
        self.status_var.set("Settings saved")
    
    def show_settings(self):
        # Switch to settings tab
        self.notebook.select(1)
    
    def show_about(self):
        about_dialog = tk.Toplevel(self.root)
        about_dialog.title("About Inventory Slip Generator")
        about_dialog.geometry("400x300")
        about_dialog.resizable(False, False)
        about_dialog.transient(self.root)
        about_dialog.grab_set()
        
        # Make dialog modal
        about_dialog.focus_set()
        
        # Configure about dialog
        about_dialog.configure(bg=self.colors.get("bg_main"))
        
        # About content
        about_frame = ttk.Frame(about_dialog, style="TFrame", padding=(20, 20))
        about_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(
            about_frame,
            text="Inventory Slip Generator",
            font=("Arial", 16, "bold"),
            style="TLabel"
        ).pack(pady=(0, 5))
        
        ttk.Label(
            about_frame,
            text=f"Version {APP_VERSION}",
            style="TLabel"
        ).pack(pady=(0, 20))
        
        ttk.Label(
            about_frame,
            text="An application for generating inventory slips from CSV and JSON data.\n"
                 "Supports Bamboo Transfer schema for cannabis inventory management.",
            justify=tk.CENTER,
            wraplength=360,
            style="TLabel"
        ).pack(pady=(0, 20))
        
        ttk.Label(
            about_frame,
            text="Created by Adam Cordova\nAGT Bothell 2025©",
            justify=tk.CENTER,
            style="TLabel"
        ).pack(pady=(0, 20))
        
        ttk.Button(
            about_frame,
            text="Close",
            command=about_dialog.destroy,
            style="TButton"
        ).pack()
    
    def show_url_dialog(self):
        url_dialog = tk.Toplevel(self.root)
        url_dialog.title("Load JSON from URL")
        url_dialog.geometry("500x120")
        url_dialog.resizable(False, False)
        url_dialog.transient(self.root)
        url_dialog.grab_set()
        
        # Make dialog modal
        url_dialog.focus_set()
        
        # Configure url dialog
        url_dialog.configure(bg=self.colors.get("bg_main"))
        
        # URL entry frame
        url_frame = ttk.Frame(url_dialog, style="TFrame", padding=(10, 10))
        url_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(url_frame, text="Enter JSON URL:", style="TLabel").pack(anchor=tk.W)
        
        url_var = tk.StringVar()
        if self.recent_urls:
            url_var.set(self.recent_urls[0])
        
        url_entry = tk.Entry(
            url_frame,
            textvariable=url_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main")
        )
        url_entry.pack(fill=tk.X, pady=(5, 10))
        url_entry.focus_set()
        url_entry.selection_range(0, tk.END)
        
        button_frame = ttk.Frame(url_frame, style="TFrame")
        button_frame.pack(fill=tk.X)
        
        ttk.Button(
            button_frame,
            text="Load",
            command=lambda: self.load_from_url(url_var.get(), url_dialog),
            style="TButton"
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        ttk.Button(
            button_frame,
            text="Cancel",
            command=url_dialog.destroy,
            style="TButton"
        ).pack(side=tk.RIGHT)
        
        # Bind Enter key to Load command
        url_entry.bind("<Return>", lambda event: self.load_from_url(url_var.get(), url_dialog))
    
    def show_json_paste_dialog(self):
        paste_dialog = tk.Toplevel(self.root)
        paste_dialog.title("Paste JSON Data")
        paste_dialog.geometry("600x400")
        paste_dialog.resizable(True, True)
        paste_dialog.transient(self.root)
        paste_dialog.grab_set()
        
        # Make dialog modal
        paste_dialog.focus_set()
        
        # Configure paste dialog
        paste_dialog.configure(bg=self.colors.get("bg_main"))
        
        # Paste area frame
        paste_frame = ttk.Frame(paste_dialog, style="TFrame", padding=(10, 10))
        paste_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(paste_frame, text="Paste JSON data below:", style="TLabel").pack(anchor=tk.W)
        
        # Scrollbar
        paste_scroll = ttk.Scrollbar(paste_frame)
        paste_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Text widget for JSON
        paste_text = tk.Text(
            paste_frame,
            font=("Consolas", 10),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main"),
            yscrollcommand=paste_scroll.set
        )
        paste_text.pack(fill=tk.BOTH, expand=True, pady=(5, 10))
        paste_scroll.config(command=paste_text.yview)
        paste_text.focus_set()
        
        # Create context menu for text area
        self.create_context_menu(paste_text)
        
        button_frame = ttk.Frame(paste_frame, style="TFrame")
        button_frame.pack(fill=tk.X)
        
        ttk.Button(
            button_frame,
            text="Import",
            command=lambda: self.process_json_data(paste_text.get(1.0, tk.END), paste_dialog),
            style="TButton"
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        ttk.Button(
            button_frame,
            text="Cancel",
            command=paste_dialog.destroy,
            style="TButton"
        ).pack(side=tk.RIGHT)
    
    def load_json(self):
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showerror("Error", "Please enter a URL.")
            return
        
        # Check if it's a Bamboo URL
        if "bamboo" in url.lower() or "getbamboo" in url.lower():
            self.auto_fetch_from_bamboo(url)
        else:
            self.load_from_url(url)
    
    def load_from_url(self, url, dialog=None):
        if not url.startswith("http"):
            messagebox.showerror("Error", "Please enter a valid URL starting with http:// or https://.")
            return
        
        self.status_var.set(f"Loading data from {url}...")
        self.progress_var.set(10)
        
        def fetch_data():
            try:
                with urllib.request.urlopen(url) as resp:
                    data = json.loads(resp.read().decode())
                
                self.root.after(0, lambda: self.process_json_data(data, dialog))
                
                # Add to recent URLs if not already there
                if url not in self.recent_urls:
                    self.recent_urls.insert(0, url)
                    self.recent_urls = self.recent_urls[:10]  # Keep only 10 most recent
                    self.config['PATHS']['recent_urls'] = '|'.join(self.recent_urls)
                    save_config(self.config)
                    self.update_recent_menu()
            
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to fetch JSON:\n{e}"))
                self.root.after(0, lambda: self.status_var.set("Failed to load data."))
                self.root.after(0, lambda: self.progress_var.set(0))
        
        # Run in a separate thread to prevent UI freezing
        threading.Thread(target=fetch_data, daemon=True).start()
    
    def process_json_data(self, data, dialog=None):
        if dialog:
            dialog.destroy()
        
        self.status_var.set("Processing JSON data...")
        self.progress_var.set(30)
        
        try:
            # If data is a string, parse it
            if isinstance(data, str):
                try:
                    data = json.loads(data)
                except json.JSONDecodeError:
                    messagebox.showerror("Error", "Invalid JSON format. Please check your data.")
                    self.status_var.set("Failed to process data.")
                    self.progress_var.set(0)
                    return
            
            # Detect and parse the format
            result_df, format_type = parse_inventory_json(data)
            
            if result_df is None:
                messagebox.showerror("Error", f"Could not parse data: {format_type}")
                self.status_var.set("Failed to process data.")
                self.progress_var.set(0)
                return
            
            # Success
            self.df = result_df
            
            # Update JSON text area
            if format_type in ["Bamboo", "Cultivera"]:
                self.json_text.delete(1.0, tk.END)
                self.json_text.insert(tk.END, json.dumps(data, indent=2))
                
                # Update tab name to reflect the format
                self.notebook.tab(2, text=f"{format_type} Import")
            
            # Sort the DataFrame for better organization
            if not self.df.empty:
                try:
                    sort_cols = []
                    if "Product Type*" in self.df.columns:
                        sort_cols.append("Product Type*")
                    if "Product Name*" in self.df.columns:
                        sort_cols.append("Product Name*")
                    
                    if sort_cols:
                        self.df = self.df.sort_values(sort_cols, ascending=[True, True])
                except:
                    pass  # If sorting fails, continue without sorting
            
            self.refresh_product_list()
            self.status_var.set(f"{format_type} data processed successfully.")
            self.progress_var.set(100)
            
            # Switch to Data tab
            self.notebook.select(0)
            
            # Reset progress after a delay
            self.root.after(2000, lambda: self.progress_var.set(0))
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process data:\n{e}")
            self.status_var.set("Failed to process data.")
            self.progress_var.set(0)
    
    def load_csv(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if not file_path:
            return
        
        self.load_csv_from_path(file_path)
    
    def load_csv_from_path(self, file_path):
        self.status_var.set(f"Loading data from {os.path.basename(file_path)}...")
        self.progress_var.set(10)
        
        def load_data():
            try:
                df = pd.read_csv(file_path)
                
                # Add to recent files if not already there
                if file_path not in self.recent_files:
                    self.recent_files.insert(0, file_path)
                    self.recent_files = self.recent_files[:10]  # Keep only 10 most recent
                    self.config['PATHS']['recent_files'] = '|'.join(self.recent_files)
                    save_config(self.config)
                    self.update_recent_menu()
                
                self.root.after(0, lambda: self.process_csv_data(df))
            
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to load CSV:\n{e}"))
                self.root.after(0, lambda: self.status_var.set("Failed to load data."))
                self.root.after(0, lambda: self.progress_var.set(0))
        
        # Run in a separate thread to prevent UI freezing
        threading.Thread(target=load_data, daemon=True).start()
    
    def process_csv_data(self, df):
        self.status_var.set("Processing CSV data...")
        self.progress_var.set(50)
        
        try:
            # Map column names to expected format
            col_map = {
                "Product Name*": "Product Name*",
                "Product Name": "Product Name*",
                "Quantity Received": "Quantity Received*",
                "Quantity*": "Quantity Received*",
                "Quantity": "Quantity Received*",
                "Lot Number*": "Barcode*",
                "Barcode": "Barcode*",
                "Lot Number": "Barcode*",
                "Accepted Date": "Accepted Date",
                "Vendor": "Vendor",
                "Strain Name": "Strain Name",
                "Product Type*": "Product Type*",
                "Product Type": "Product Type*",
                "Inventory Type": "Product Type*"
            }
            
            df = df.rename(columns=lambda x: col_map.get(x.strip(), x.strip()))
            
            # Ensure required columns exist
            required_cols = ["Product Name*", "Barcode*"]
            missing_cols = [col for col in required_cols if col not in df.columns]
            
            if missing_cols:
                messagebox.showerror("Error", f"CSV is missing required columns: {', '.join(missing_cols)}")
                self.status_var.set("Failed to process data.")
                self.progress_var.set(0)
                return
            
            # Set default values for missing columns
            if "Vendor" not in df.columns:
                df["Vendor"] = "Unknown Vendor"
            else:
                df["Vendor"] = df["Vendor"].fillna("Unknown Vendor")
            
            if "Accepted Date" not in df.columns:
                today = datetime.datetime.today().strftime("%Y-%m-%d")
                df["Accepted Date"] = today
            
            if "Product Type*" not in df.columns:
                df["Product Type*"] = "Unknown"
            
            if "Strain Name" not in df.columns:
                df["Strain Name"] = ""
            
            # Store the DataFrame
            self.df = df
            
            # Sort if possible
            try:
                sort_cols = []
                if "Product Type*" in self.df.columns:
                    sort_cols.append("Product Type*")
                if "Product Name*" in self.df.columns:
                    sort_cols.append("Product Name*")
                
                if sort_cols:
                    self.df = self.df.sort_values(sort_cols, ascending=[True, True])
            except:
                pass  # If sorting fails, continue without sorting
            
            self.refresh_product_list()
            self.status_var.set("CSV data processed successfully.")
            self.progress_var.set(100)
            
            # Reset progress after a delay
            self.root.after(2000, lambda: self.progress_var.set(0))
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process CSV data:\n{e}")
            self.status_var.set("Failed to process data.")
            self.progress_var.set(0)
    
    def load_bamboo_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
        if not file_path:
            return
        
        self.status_var.set(f"Loading Bamboo data from {os.path.basename(file_path)}...")
        self.progress_var.set(10)
        
        def load_data():
            try:
                with open(file_path, 'r') as f:
                    data = json.load(f)
                
                # Add to recent files if not already there
                if file_path not in self.recent_files:
                    self.recent_files.insert(0, file_path)
                    self.recent_files = self.recent_files[:10]  # Keep only 10 most recent
                    self.config['PATHS']['recent_files'] = '|'.join(self.recent_files)
                    save_config(self.config)
                    self.update_recent_menu()
                
                # Update JSON text area
                self.root.after(0, lambda: self.json_text.delete(1.0, tk.END))
                self.root.after(0, lambda: self.json_text.insert(tk.END, json.dumps(data, indent=2)))
                
                # Process the data
                self.root.after(0, lambda: self.process_json_data(data))
            
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to load JSON file:\n{e}"))
                self.root.after(0, lambda: self.status_var.set("Failed to load data."))
                self.root.after(0, lambda: self.progress_var.set(0))
        
        # Run in a separate thread to prevent UI freezing
        threading.Thread(target=load_data, daemon=True).start()
    
    def import_bamboo_data(self):
        # Get JSON data from text area
        json_data = self.json_text.get(1.0, tk.END)
        if not json_data.strip():
            messagebox.showwarning("Warning", "No JSON data to import.")
            return
        
        try:
            self.process_json_data(json_data)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to import Bamboo data:\n{e}")
    
    def clear_json_data(self):
        self.json_text.delete(1.0, tk.END)
    
    def change_theme(self, theme_name):
        if self.colors.switch_theme(theme_name):
            self.theme_name = theme_name
            self.config['SETTINGS']['theme'] = theme_name
            save_config(self.config)
            
            # Show message about need to restart
            messagebox.showinfo("Theme Changed", 
                             "Theme changes will be fully applied after restarting the application.")
    
    def refresh_product_list(self):
        # Clear existing product list
        for widget in self.product_frame.winfo_children():
            widget.destroy()
        
        self.product_vars.clear()
        self.group_vars.clear()
        
        # If no data, show empty message
        if self.df.empty:
            self.empty_label = ttk.Label(
                self.product_frame,
                text="No products loaded. Please load data from CSV or JSON.",
                style="TLabel",
                font=("Arial", 12)
            )
            self.empty_label.pack(pady=50)
            return
        
        # Group data by product type
        grouped = self.df.groupby("Product Type*") if "Product Type*" in self.df.columns else {"All Products": self.df}
        
        # For each product type, create a group
        for product_type, group_df in grouped:
            # Add group select checkbox
            group_var = tk.BooleanVar(value=True)
            self.group_vars[product_type] = group_var
            
            # Create frame for group header
            group_frame = ttk.Frame(self.product_frame, style="TFrame")
            group_frame.pack(fill=tk.X, pady=(10, 5), padx=5)
            
            # Group checkbox
            group_cb = ttk.Checkbutton(
                group_frame,
                text=f"Select/Deselect {product_type}",
                variable=group_var,
                command=lambda pt=product_type: self.toggle_group(pt),
                style="TCheckbutton"
            )
            group_cb.pack(fill=tk.X, pady=2)
            
            # Add each product under the group
            for idx, row in group_df.iterrows():
                product_name = row.get("Product Name*", "")
                sku = str(row.get("Barcode*", "")).strip()
                strain_name = row.get("Strain Name", "")
                
                qty = ""
                if "Quantity Received*" in row and pd.notna(row["Quantity Received*"]):
                    qty = row["Quantity Received*"]
                elif "Quantity*" in row and pd.notna(row["Quantity*"]):
                    qty = row["Quantity*"]
                
                if pd.notna(qty):
                    try:
                        qty = int(float(qty))
                    except (ValueError, TypeError):
                        qty = str(qty)
                else:
                    qty = ""
                
                # Create row with product info
                row_frame = ttk.Frame(self.product_frame, style="TFrame")
                row_frame.pack(fill=tk.X, pady=2, padx=10)
                
                # Store metadata for search functionality
                row_frame.product_name = product_name
                row_frame.product_type = product_type
                row_frame.strain_name = strain_name
                
                # Checkbox for selection
                var = tk.BooleanVar(value=True)
                cb = ttk.Checkbutton(
                    row_frame,
                    text=f"{product_name}",
                    variable=var,
                    style="TCheckbutton"
                )
                cb.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                # SKU and Quantity labels
                info_frame = ttk.Frame(row_frame, style="TFrame")
                info_frame.pack(side=tk.RIGHT)
                
                if strain_name:
                    ttk.Label(
                        info_frame,
                        text=f"Strain: {strain_name}",
                        style="TLabel",
                        font=("Arial", 9)
                    ).pack(side=tk.TOP, anchor=tk.E)
                
                ttk.Label(
                    info_frame,
                    text=f"SKU: {sku}",
                    style="TLabel",
                    font=("Arial", 9)
                ).pack(side=tk.TOP, anchor=tk.E)
                
                ttk.Label(
                    info_frame,
                    text=f"Qty: {qty}",
                    style="TLabel",
                    font=("Arial", 10, "bold")
                ).pack(side=tk.TOP, anchor=tk.E)
                
                # Store variable and product type for bulk selection
                self.product_vars[idx] = (var, product_type)
    
    def on_generate(self):
        # Get selected products
        sel_idxs = [i for i, (var, _) in self.product_vars.items() if var.get()]
        
        if not sel_idxs:
            messagebox.showerror("Error", "No products selected.")
            return
        
        selected_df = self.df.loc[sel_idxs].copy()
        
        # Get latest settings from UI
        self.config['PATHS']['output_dir'] = self.output_dir_var.get()
        self.config['PATHS']['template_path'] = self.template_path_var.get()
        self.config['SETTINGS']['items_per_page'] = self.items_per_page_var.get()
        self.config['SETTINGS']['auto_open'] = str(self.auto_open_var.get())
        save_config(self.config)
        
        # Update status
        self.status_var.set("Generating inventory slips...")
        self.progress_var.set(0)
        
        # Run in a separate thread
        def generate():
            success, result = run_full_process_inventory_slips(
                selected_df,
                self.config,
                lambda msg: self.root.after(0, lambda: self.status_var.set(msg)),
                lambda val: self.root.after(0, lambda: self.progress_var.set(val))
            )
            
            if success:
                self.root.after(0, lambda: messagebox.showinfo("Success", f"Inventory slips saved to:\n{result}"))
            else:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to generate inventory slips:\n{result}"))
            
            # Reset progress after a delay
            self.root.after(2000, lambda: self.progress_var.set(0))
        
        threading.Thread(target=generate, daemon=True).start()
        def auto_fetch_from_bamboo(self, url=None):
            """Handle Bamboo API access with proper authentication or fallback to manual data import"""
            if not url:
                url = self.url_entry.get().strip()
                if not url:
                    messagebox.showerror("Error", "Please enter a Bamboo API URL.")
                    return
            
            self.status_var.set("Connecting to Bamboo API...")
            self.progress_var.set(10)
            
            # Setup headers for Bamboo API
            headers = {
                "User-Agent": "InventorySlipGenerator/2.0.0",
                "Accept": "application/json",
                "Content-Type": "application/json"
            }
            
            # Check if we have API key in config
            if 'API' not in self.config:
                self.config['API'] = {}
            
            api_key = self.config['API'].get('bamboo_key', '')
            if api_key:
                headers["Authorization"] = f"Bearer {api_key}"
            
            def fetch_data():
                try:
                    req = urllib.request.Request(url, headers=headers)
                    with urllib.request.urlopen(req) as resp:
                        data = json.loads(resp.read().decode())
                        
                        # Process the data
                        self.root.after(0, lambda: self.process_json_data(data))
                        
                        # Cache the response for offline use
                        cache_dir = os.path.join(os.path.expanduser("~"), ".inventory_slip_cache")
                        if not os.path.exists(cache_dir):
                            os.makedirs(cache_dir)
                        
                        cache_file = os.path.join(cache_dir, "bamboo_latest.json")
                        with open(cache_file, 'w') as f:
                            json.dump(data, f)
                        
                        # Add to recent URLs
                        if url not in self.recent_urls:
                            self.recent_urls.insert(0, url)
                            self.recent_urls = self.recent_urls[:10]
                            self.config['PATHS']['recent_urls'] = '|'.join(self.recent_urls)
                            save_config(self.config)
                            self.update_recent_menu()
                
                except urllib.error.HTTPError as e:
                    if e.code == 403:
                        # Handle "forbidden" error - try to use cached data
                        self.root.after(0, lambda: self.handle_bamboo_forbidden())
                    else:
                        self.root.after(0, lambda: messagebox.showerror("Error", f"API Error: {e.code} - {e.reason}"))
                        self.root.after(0, lambda: self.status_var.set(f"Failed to fetch data: {e.reason}"))
                        self.root.after(0, lambda: self.progress_var.set(0))
                
                except Exception as e:
                    self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to fetch data: {str(e)}"))
                    self.root.after(0, lambda: self.status_var.set("Failed to fetch data."))
                    self.root.after(0, lambda: self.progress_var.set(0))
            
            # Run in a separate thread
            threading.Thread(target=fetch_data, daemon=True).start()

    def handle_bamboo_forbidden(self):
        """Handle case where Bamboo access is forbidden"""
        self.status_var.set("Bamboo API access forbidden - checking cache...")
        
        # Check if we have cached data
        cache_file = os.path.join(os.path.expanduser("~"), ".inventory_slip_cache", "bamboo_latest.json")
        
        if os.path.exists(cache_file):
            try:
                # Use cached data instead
                with open(cache_file, 'r') as f:
                    data = json.load(f)
                
                # Update JSON text area with cached data
                self.json_text.delete(1.0, tk.END)
                self.json_text.insert(tk.END, json.dumps(data, indent=2))
                
                # Process the cached data
                self.process_json_data(data)
                
                # Notify user
                self.status_var.set("Using cached Bamboo data (API access forbidden)")
                messagebox.showinfo("Access Restricted", 
                                "Bamboo API access is forbidden. Using cached data instead.\n\n"
                                "Please verify your API credentials in Settings.")
                
                # Add API Settings button to manage credentials
                self.show_api_settings()
            except Exception as e:
                self.show_json_paste_dialog()
        else:
            self.show_json_paste_dialog()

    def show_api_settings(self):
        """Show dialog to configure API settings"""
        settings_dialog = tk.Toplevel(self.root)
        settings_dialog.title("API Settings")
        settings_dialog.geometry("500x250")
        settings_dialog.resizable(False, False)
        settings_dialog.transient(self.root)
        settings_dialog.grab_set()
        
        # Make dialog modal
        settings_dialog.focus_set()
        
        # Configure dialog
        settings_dialog.configure(bg=self.colors.get("bg_main"))
        
        # Settings frame
        settings_frame = ttk.Frame(settings_dialog, style="TFrame", padding=(10, 10))
        settings_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        ttk.Label(
            settings_frame,
            text="API Integration Settings",
            font=("Arial", 14, "bold"),
            style="TLabel"
        ).pack(pady=(0, 20))
        
        # Ensure API section exists in config
        if 'API' not in self.config:
            self.config['API'] = {}
        
        # Bamboo section
        bamboo_frame = ttk.LabelFrame(settings_frame, text="Bamboo API", style="TFrame")
        bamboo_frame.pack(fill=tk.X, pady=(0, 10))
        
        api_key_frame = ttk.Frame(bamboo_frame, style="TFrame")
        api_key_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(
            api_key_frame,
            text="API Key:",
            style="TLabel"
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        bamboo_key_var = tk.StringVar(value=self.config['API'].get('bamboo_key', ''))
        
        api_key_entry = tk.Entry(
            api_key_frame,
            textvariable=bamboo_key_var,
            font=("Arial", 11),
            bg=self.colors.get("entry_bg"),
            fg=self.colors.get("entry_fg"),
            insertbackground=self.colors.get("fg_main"),
            show="•"  # Hide the API key for security
        )
        api_key_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Instructions
        ttk.Label(
            bamboo_frame,
            text="Enter your Bamboo API key to enable direct access to inventory data.\n"
                "You can find this in your Bamboo account settings.",
            style="TLabel",
            wraplength=450
        ).pack(padx=10, pady=(0, 10))
        
        # Button frame
        button_frame = ttk.Frame(settings_frame, style="TFrame")
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        ttk.Button(
            button_frame,
            text="Save",
            command=lambda: self.save_api_settings(bamboo_key_var.get(), settings_dialog),
            style="TButton"
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        ttk.Button(
            button_frame,
            text="Cancel",
            command=settings_dialog.destroy,
            style="TButton"
        ).pack(side=tk.RIGHT)

    def save_api_settings(self, bamboo_key, dialog):
        """Save API settings to config"""
        if 'API' not in self.config:
            self.config['API'] = {}
        
        self.config['API']['bamboo_key'] = bamboo_key
        save_config(self.config)
        
        # Show success message
        self.status_var.set("API settings saved successfully.")
        dialog.destroy()
    
    def on_close(self):
        # Save settings before closing
        save_config(self.config)
        self.root.destroy()


def main():
    root = tk.Tk()
    
    # Set window title and icon
    root.title(f"Inventory Slip Generator v{APP_VERSION}")
    
    try:
        if sys.platform == "win32":
            root.iconbitmap(resource_path("assets/icon.ico"))
    except:
        pass
    
    # Center window on screen
    app_width = 850
    app_height = 750
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = (screen_width // 2) - (app_width // 2)
    y = (screen_height // 2) - (app_height // 2)
    root.geometry(f"{app_width}x{app_height}+{x}+{y}")
    
    # Create application
    app = InventorySlipGenerator(root)
    
    # Start the main loop
    root.mainloop()


if __name__ == "__main__":
    main()